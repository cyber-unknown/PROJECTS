#importing the required modules .....

import os.path
import tkinter
import customtkinter
from tkinter import Tk, Canvas, Entry, Button, PhotoImage,filedialog
from PIL import Image
import aspose.words as aw
import docx
import PyPDF2
from pathlib import Path
#defining functions.....

def convert_btn():
    if btn == 1:
        pdftoimage()
    elif btn == 2:
        pdftoword()
    elif btn == 3:
        wordtopdf()                  #convert_btn defines the function of convert button
    elif btn == 4:
        wordtoimage()
    elif btn == 5:
        imagetopdf()
    elif btn == 6:
        imagetoword()

def inputpath():
    entry_1.delete(0,'end')
    filepath = filedialog.askopenfilename()      #this function is defined to get the inputfile
    entry_1.insert(tkinter.END,filepath)
    global path
    path=filepath

def pdftoimage():
    from pdf2image import convert_from_path
    images = convert_from_path(path, 500, poppler_path=r'C:\Program Files\poppler-21.11.0\Library\bin',)
    for i, image in enumerate(images):
        fname = 'image' + str(i) + '.png'
        image.save(fname, "PNG")

def pdftoword():
    with open(path, mode='rb') as f:
        reader = PyPDF2.PdfFileReader(f)
        page = reader.getPage(0)
        print(page.extractText())          #this function is defined to convert pdf_to_word
    g = page.extractText()
    d=path.replace('pdf','doc')
    h = open(d,'w')
    h.write(g)
    h.close()

def wordtopdf():
    from docx2pdf import convert
    e = path.replace('docx', 'pdf')        #this function is defined to convert word_to_pdf
    convert(path, e)

def wordtoimage():
    doc = aw.Document(path)
    shapes = doc.get_child_nodes(aw.NodeType.SHAPE, True)
    imageIndex = 0
    for shape in shapes:
        shape = shape.as_shape()                #this function is defined to convert word_to_image
        if (shape.has_image):
            c= os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            imageFileName = c+'\\'+ f"exported image.{imageIndex}_{aw.FileFormatUtil.image_type_to_extension(shape.image_data.image_type)}"

            print(imageFileName)
            shape.image_data.save(imageFileName)
            imageIndex += 1
def imagetopdf():
    image1 = Image.open(path)
    im1 = image1.convert('RGB')
    c='pdf'
    if path[-1:-4:-1] == 'gep':
        d = path.replace('jpeg', c)         #this function is defined to convert image_to_pdf
    elif path[-1:-4:-1] == 'gpj':
        d = path.replace('jpg', c)
    else:
        d = path.replace('png', c)
    im1.save(d)

def imagetoword():
    c='docx'
    if path[-1:-4:-1] == 'gep':
        d=path.replace('jpeg',c)
    elif path[-1:-4:-1] == 'gpj':     #this function is defined to convert image_to_word
        d=path.replace('jpg',c)
    else:
        d=path.replace('png',c)
    doc = docx.Document()
    doc.add_picture(path)
    doc.save(d)

def slider_changed(val):
    a=slider.get()
    if a < 2/2:
        global btn
        btn=1
    elif a > 2/2 and a < 4/2:
        btn=2
    elif a > 4/2 and a < 6/2:    #this function is defined to get the slider value and assign btn value
        btn=3
    elif a > 6/2 and a <8/2:
        btn=4
    elif a > 8/2 and a <10/2:
        btn=5
    elif a >10/2 and a <12/2:
        btn=6

OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path("./assets")


def relative_to_assets(path: str) -> Path:   #this function is defined to get the image_source of gui
    return ASSETS_PATH / Path(path)

window = Tk()                         #Creates root master with the TK()
window.geometry("729x452")
window.configure(bg = "#000000")       #creates a tkinter window
window.title('FILE CONVERSION IO')

canvas = Canvas(
    window,
    bg = "#000000",
    height = 452,
    width = 729,                       #creates the background of the tkinter window
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)

canvas.place(x = 0, y = 0)
canvas.create_text(
    245.0,
    19.0,
    anchor="nw",
    text="FILE CONVERSION I/O",          #creates title of the program
    fill="#DBFF00",
    font=("Roboto", 24 * -1)
)

button_image_1 = PhotoImage(
    file=relative_to_assets("button_1.png"))     #retrives button_1's image from asserts folder
button_1 = Button(
    image=button_image_1,
    borderwidth=0,                               #created button that uses convert_btn function
    highlightthickness=0,
    command=lambda: convert_btn(),
    relief="flat"
)

button_1.place(
    x=270.0,
    y=334.0,                                   #places the convert button
    width=190.0,
    height=39.0
)

button_image_2 = PhotoImage(
    file=relative_to_assets("button_2.png"))  #retrives the button_2's image form asserts folder
button_2 = Button(master=window,
    image=button_image_2,
    borderwidth=0,
    highlightthickness=0,            #creates button that uses inputpath function
    command=lambda: inputpath(),
    relief="flat"
)

button_2.place(
    x=36.0,
    y=93.0,                          #places the select file button
    width=107.0,
    height=51.0
)

entry_image_1 = PhotoImage(
    file=relative_to_assets("entry_1.png"))     #retrives the entry_1's image from asserts folder
entry_bg_1 = canvas.create_image(
    408.5,
    115.5,                                      #create the background for entry_1
    image=entry_image_1
)

entry_1 = Entry(master=window,
    bd=0,
    bg="#C0C0BB",                           #creates a entry to display path of inputfile
    highlightthickness=0
)

entry_1.place(
    x=179.0,
    y=104.0,                                 #places entry_1
    width=459.0,
    height=27.0
)

image_image_1 = PhotoImage(
    file=relative_to_assets("image_1.png"))   #retrives the python logo from assertes folder
image_1 = canvas.create_image(
    50.0,                                     #displayes the python logo
    41.0,
    image=image_image_1
)

image_image_2 = PhotoImage(
    file=relative_to_assets("image_2.png"))  #retrives image of the arrow markings above the slider
image_2 = canvas.create_image(
    365.0,
    222.0,                                   #displays the arrow markings
    image=image_image_2
)

slider = customtkinter.CTkSlider(master=window,
                                 width=478,
                                 height=30,
                                 border_width=5.5,    #creates a slider
                                 from_=0,
                                 to=6,
                                 command=slider_changed)
slider.place(relx=0.5, rely=0.6, anchor=tkinter.CENTER)   #places the slider

window.resizable(False, False)  #makes tkinter window to stay in same resolution

window.mainloop()           #main event loop
